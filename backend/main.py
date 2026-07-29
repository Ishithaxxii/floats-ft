from fastapi import FastAPI, HTTPException, Query, Header, Depends, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

import os
import re
import sqlite3
import pandas as pd
import smtplib
import uuid
import json
from email.message import EmailMessage

from meta_processor import infer_ship_from_metadata

from datetime import date
from typing import Optional

# ==========================================
# FASTAPI SETUP
# ==========================================

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# DIRECTORIES
# ==========================================

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(BASE_DIR)
DATA_DIR    = os.path.join(PROJECT_ROOT, "data")
CACHE_DIR   = os.path.join(BASE_DIR, "cache")

CACHE_DB = os.environ.get(
    "SEASNAP_CACHE_DB",
    os.path.join(CACHE_DIR, "cache_profiles.sqlite"),
)

# ==========================================
# API KEY AUTH
# ==========================================
from dotenv import load_dotenv
load_dotenv()

API_KEY = os.environ.get("SEASNAP_API_KEY", "52222a785cb7ffbe8223d02b9ec90f5cff04eaff1b055db430c9d53bd4f8d90c")

def require_api_key(x_api_key: str = Header(None)):
    if x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key")

# ==========================================
# EEZ BOUNDARY (server-side check — source of truth)
# ==========================================
# Reuses the SAME point-sequence GeoJSON the frontend already loads for
# drawing the EEZ boundary line (india_eez.geojson: a sequence of Point
# features, broken into loops wherever consecutive points jump too far
# apart). We port the frontend's extractEEZLoops() + ray-casting
# point-in-polygon check to Python, so both sides agree on what "inside
# the EEZ" means without needing a second polygon file.
#
# IMPORTANT: update SEASNAP_EEZ_GEOJSON (or the default path below) to
# point at wherever india_eez.geojson actually lives in your repo.

EEZ_GEOJSON_PATH = os.environ.get(
    "SEASNAP_EEZ_GEOJSON",
    "/home/ishitha/Desktop/seasnap-ft/frontend/public/data/india_eez.geojson",
)

EEZ_BREAK_THRESHOLD = 3.0  # matches BREAK_THRESHOLD in MapView.jsx's extractEEZLoops

_eez_loops: list[list[tuple[float, float]]] = []  # each loop: list of (lat, lon)


def _extract_eez_loops(geojson: dict) -> list[list[tuple[float, float]]]:
    points = [
        (feat["geometry"]["coordinates"][1], feat["geometry"]["coordinates"][0])
        for feat in geojson.get("features", [])
        if feat.get("geometry", {}).get("type") == "Point"
    ]
    if not points:
        return []

    loops: list[list[tuple[float, float]]] = []
    current_loop = [points[0]]

    for i in range(1, len(points)):
        prev_lat, prev_lon = points[i - 1]
        curr_lat, curr_lon = points[i]
        distance = ((curr_lat - prev_lat) ** 2 + (curr_lon - prev_lon) ** 2) ** 0.5
        if distance > EEZ_BREAK_THRESHOLD:
            if len(current_loop) > 1:
                loops.append(current_loop)
            current_loop = [points[i]]
        else:
            current_loop.append(points[i])

    if len(current_loop) > 1:
        loops.append(current_loop)

    return loops


def _load_eez_loops() -> None:
    global _eez_loops
    if not os.path.isfile(EEZ_GEOJSON_PATH):
        print(
            f"[eez] WARNING: {EEZ_GEOJSON_PATH} not found — "
            f"EEZ restriction will NOT be enforced! Set SEASNAP_EEZ_GEOJSON."
        )
        _eez_loops = []
        return
    with open(EEZ_GEOJSON_PATH) as f:
        geojson = json.load(f)
    _eez_loops = _extract_eez_loops(geojson)
    total_pts = sum(len(loop) for loop in _eez_loops)
    print(f"[eez] Loaded {len(_eez_loops)} EEZ loop(s), {total_pts} points, from {EEZ_GEOJSON_PATH}")


def _ray_cast_point_in_polygon(lat: float, lon: float, polygon: list[tuple[float, float]]) -> bool:
    inside = False
    n = len(polygon)
    j = n - 1
    for i in range(n):
        yi, xi = polygon[i]
        yj, xj = polygon[j]
        if (yi > lat) != (yj > lat):
            x_intersect = (xj - xi) * (lat - yi) / (yj - yi) + xi
            if lon < x_intersect:
                inside = not inside
        j = i
    return inside


def is_in_eez(lat: float, lon: float) -> bool:
    return any(_ray_cast_point_in_polygon(lat, lon, loop) for loop in _eez_loops)


_load_eez_loops()

# ==========================================
# INSTRUMENT CONFIG
# ==========================================
# column_map  : raw CSV column -> canonical name
# output_columns : canonical columns stored in SQLite and returned by /profile

INSTRUMENT_CONFIG: dict[str, dict] = {
    "ctd": {
        "data_folder": os.environ.get(
            "SEASNAP_CTD_DATA_FOLDER",
            "/home/ishitha/CTD",
        ),
        "meta_folder": os.environ.get(
            "SEASNAP_CTD_META_FOLDER",
            "/home/ishitha/CTD/metadata",
        ),
        "table": "profiles_ctd",
        "column_map": {
            "Depth (m)":           "depSM",
            "depSM":               "depSM",
            "Temp-90 (deg C)":     "TEMP_QC_VAR",
            "t090C":               "TEMP_QC_VAR",
            "TEMP_QC_VAR":         "TEMP_QC_VAR",
            "Sal00 (psu)":         "SAL_QC_VAR",
            "Sal00":               "SAL_QC_VAR",
            "SAL_QC_VAR":          "SAL_QC_VAR",
            "Conductivity (S/m)":  "c0S/m",
            "c0S/m":               "c0S/m",
            "Sigma-t":             "sigma-t00",
            "sigma-t00":           "sigma-t00",
            "DO (ml/l)":           "sbeox0ML/L",
            "sbeox0ML/L":          "sbeox0ML/L",
            "SourceFile":          "SourceFile",
            "folderpath_filename": "SourceFile",
            "Temp_QC":             "Temp_QC",
            "Sal_QC":              "Sal_QC",
            "Pres_QC":             "Pres_QC",
            "ALL_TESTS_QC":        "ALL_TESTS_QC",
        },
        "output_columns": [
            "depSM", "TEMP_QC_VAR", "Temp_QC",
            "SAL_QC_VAR", "Sal_QC", "c0S/m",
            "sigma-t00", "sbeox0ML/L", "Pres_QC", "ALL_TESTS_QC",
        ],
    },

    "xbt": {
        "data_folder": os.environ.get(
            "SEASNAP_XBT_DATA_FOLDER",
            "/home/ishitha/XBT",
        ),
        "meta_folder": os.environ.get(
            "SEASNAP_XBT_META_FOLDER",
            "/home/ishitha/XBT/metadata",
        ),
        "table": "profiles_xbt",
        "column_map": {
            "Depth (m)":           "depSM",
            "Temperature (deg C)": "TEMP_QC_VAR",
            "TEMP_QC_VAR":         "TEMP_QC_VAR",
            "Temp_QC":             "Temp_QC",
            "Pres_QC":             "Pres_QC",
            "SourceFile":          "SourceFile",
            "ALL_TESTS_QC":        "ALL_TESTS_QC",
        },
        "output_columns": [
            "depSM", "TEMP_QC_VAR", "Temp_QC", "Pres_QC", "ALL_TESTS_QC",
        ],
    },

    "xctd": {
        "data_folder": os.environ.get(
            "SEASNAP_XCTD_DATA_FOLDER",
            "/home/ishitha/XCTD",
        ),
        "meta_folder": os.environ.get(
            "SEASNAP_XCTD_META_FOLDER",
            "/home/ishitha/XCTD/metadata",
        ),
        "table": "profiles_xctd",
        "column_map": {
            "Depth (m)":           "depSM",
            "Temperature (deg C)": "TEMP_QC_VAR",
            "TEMP_QC_VAR":         "TEMP_QC_VAR",
            "Salinity (psu)":      "SAL_QC_VAR",
            "SAL_QC_VAR":          "SAL_QC_VAR",
            "Temp_QC":             "Temp_QC",
            "Sal_QC":              "Sal_QC",
            "Pres_QC":             "Pres_QC",
            "SourceFile":          "SourceFile",
            "ALL_TESTS_QC":        "ALL_TESTS_QC",
        },
        "output_columns": [
            "depSM", "TEMP_QC_VAR", "Temp_QC",
            "SAL_QC_VAR", "Sal_QC", "Pres_QC", "ALL_TESTS_QC",
        ],
    },
}

# Full union of columns — /profile always returns this shape;
# columns not produced by a given instrument are filled with None.
ALL_OUTPUT_COLUMNS = [
    "depSM", "TEMP_QC_VAR", "Temp_QC",
    "SAL_QC_VAR", "Sal_QC", "c0S/m",
    "sigma-t00", "sbeox0ML/L", "Pres_QC", "ALL_TESTS_QC",
]

QC_COLUMNS = {"Temp_QC", "Sal_QC", "Pres_QC", "ALL_TESTS_QC"}

# In-memory station cache: instrument_type -> list[dict]
_station_cache: dict[str, list[dict]] = {}


# ==========================================
# FOLDER HELPERS
# ==========================================

def _csv_files(folder: str) -> list[str]:
    """Sorted list of .csv paths in folder."""
    return sorted(
        os.path.join(folder, f)
        for f in os.listdir(folder)
        if f.endswith(".csv")
    )


def _resolve_folder(instrument_type: str) -> str:
    """Return a valid data folder for the instrument, or raise."""
    cfg = INSTRUMENT_CONFIG[instrument_type]
    candidates = [
        cfg["data_folder"],
        os.path.join(DATA_DIR, instrument_type),
        os.path.join(PROJECT_ROOT, instrument_type),
    ]
    for path in candidates:
        if path and os.path.isdir(path) and _csv_files(path):
            return path
    raise FileNotFoundError(
        f"{instrument_type.upper()} data folder not found or has no CSVs. "
        f"Set SEASNAP_{instrument_type.upper()}_DATA_FOLDER."
    )


def _folder_mtime(folder: str) -> float:
    """Latest mtime across all CSVs in folder."""
    files = _csv_files(folder)
    return max((os.path.getmtime(f) for f in files), default=0.0)


# ==========================================
# SQLITE CACHE HELPERS (profile data)
# ==========================================

def _meta_key(instrument_type: str, key: str) -> str:
    return f"{instrument_type}:{key}"


def _cache_is_current(instrument_type: str, folder: str) -> bool:
    if not os.path.isfile(CACHE_DB):
        return False
    latest = _folder_mtime(folder)
    try:
        with sqlite3.connect(CACHE_DB) as conn:
            rows = dict(conn.execute("SELECT key, value FROM metadata").fetchall())
    except sqlite3.OperationalError:
        return False
    return (
        rows.get(_meta_key(instrument_type, "source_path")) == folder
        and float(rows.get(_meta_key(instrument_type, "source_mtime"), 0)) == latest
    )


def _build_cache(instrument_type: str, folder: str) -> None:
    cfg = INSTRUMENT_CONFIG[instrument_type]

    table        = cfg["table"]
    column_map   = cfg["column_map"]
    out_cols     = cfg["output_columns"]
    usecols_set  = set(column_map.keys())

    csv_files    = _csv_files(folder)
    latest_mtime = _folder_mtime(folder)

    os.makedirs(CACHE_DIR, exist_ok=True)

    print(
        f"[{instrument_type}] Building cache from "
        f"{len(csv_files)} file(s) in: {folder}"
    )

    col_defs = ", ".join(
        f'"{c}" {"INTEGER" if c in QC_COLUMNS else "REAL"}'
        for c in out_cols
    )

    with sqlite3.connect(CACHE_DB) as conn:

        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=NORMAL")
        conn.execute("PRAGMA temp_store=MEMORY")
        conn.execute("PRAGMA cache_size=-100000")

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS metadata (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )
            """
        )

        conn.execute(f"DROP TABLE IF EXISTS {table}")

        conn.execute(
            f"""
            CREATE TABLE {table} (
                stem TEXT NOT NULL,
                {col_defs}
            )
            """
        )

        total_rows = 0

        for csv_path in csv_files:

            fname = os.path.basename(csv_path)

            print(f"[{instrument_type}] Reading {fname}")

            file_rows = 0

            for chunk in pd.read_csv(
                csv_path,
                usecols=lambda c: c.strip() in usecols_set,
                chunksize=200_000,
                low_memory=False,
            ):

                chunk.columns = chunk.columns.str.strip()

                chunk = chunk.rename(columns=column_map)

                chunk = chunk.loc[
                    :,
                    ~chunk.columns.duplicated(keep="last")
                ]

                if "SourceFile" not in chunk.columns:
                    raise ValueError(
                        f"No SourceFile column in {csv_path}"
                    )

                chunk["stem"] = (
                    chunk["SourceFile"]
                    .astype(str)
                    .str.rsplit(".", n=1)
                    .str[0]
                    .str.strip()
                    .str.lower()
                )

                chunk = chunk.reindex(
                    columns=["stem"] + out_cols
                )

                for col in out_cols:
                    chunk[col] = pd.to_numeric(
                        chunk[col],
                        errors="coerce"
                    )

                chunk = chunk.dropna(
                    subset=["stem", "depSM"]
                )

                rows = len(chunk)

                chunk.to_sql(
                    table,
                    conn,
                    if_exists="append",
                    index=False,
                    method="multi",
                    chunksize=5000,
                )

                file_rows += rows
                total_rows += rows

            print(
                f"  {fname}: "
                f"{file_rows:,} rows "
                f"(total {total_rows:,})"
            )

        print(
            f"[{instrument_type}] Creating stem index..."
        )

        conn.execute(
            f"CREATE INDEX IF NOT EXISTS idx_{table}_stem_dep "
            f"ON {table}(stem, depSM)"
        )

        count = conn.execute(
            f"SELECT COUNT(*) FROM {table}"
        ).fetchone()[0]

        print(
            f"[{instrument_type}] "
            f"Table rows verified: {count:,}"
        )

        conn.execute(
            "INSERT OR REPLACE INTO metadata VALUES (?, ?)",
            (_meta_key(instrument_type, "source_path"), folder),
        )

        conn.execute(
            "INSERT OR REPLACE INTO metadata VALUES (?, ?)",
            (_meta_key(instrument_type, "source_mtime"), str(latest_mtime)),
        )

    print(
        f"[{instrument_type}] "
        f"Cache ready: {total_rows:,} rows."
    )


def ensure_cache(instrument_type: str) -> str:

    folder = _resolve_folder(instrument_type)

    current = _cache_is_current(instrument_type, folder)

    print(f"[{instrument_type}] Cache current = {current}")

    if not current:
        _build_cache(instrument_type, folder)
    else:
        print(f"[{instrument_type}] Cache is current.")

    return folder

# ==========================================
# METADATA LOADING (parsing helpers)
# ==========================================

_META_RENAME = {
    "Latitude(decimal)":   "Latitude_decimal",
    "Latitude (decimal)":  "Latitude_decimal",
    "Longitude(decimal)":  "Longitude_decimal",
    "Longitude (decimal)": "Longitude_decimal",
    "Depth":               "Station Depth",
    "Station":             "Station Number",
}

_COALESCE_COLS = ["Latitude_decimal", "Longitude_decimal", "Station Depth", "Station Number"]


def _resolve_meta_folder(instrument_type: str) -> str:
    cfg = INSTRUMENT_CONFIG[instrument_type]
    candidates = [
        cfg["meta_folder"],
        os.path.join(PROJECT_ROOT, "meta", instrument_type),
        os.path.join(BASE_DIR,     "meta", instrument_type),
    ]
    for path in candidates:
        if path and os.path.isdir(path):
            return path
    return cfg["meta_folder"]  # let caller handle missing


def _load_meta_df(instrument_type: str) -> tuple[pd.DataFrame, str, str | None]:
    """Load, clean and return the metadata DataFrame for one instrument."""
    meta_folder = _resolve_meta_folder(instrument_type)

    if not os.path.isdir(meta_folder):
        return pd.DataFrame(), meta_folder, "Meta folder not found"

    csv_files = [
        os.path.join(meta_folder, f)
        for f in os.listdir(meta_folder)
        if f.endswith(".csv")
    ]
    if not csv_files:
        return pd.DataFrame(), meta_folder, "No CSV files found"

    frames = []
    for path in csv_files:
        try:
            frames.append(pd.read_csv(path))
        except Exception as e:
            print(f"[{instrument_type}] Failed to read {path}: {e}")

    if not frames:
        return pd.DataFrame(), meta_folder, "No valid CSV files"

    df = pd.concat(frames, ignore_index=True).rename(columns=_META_RENAME)

    for col in _COALESCE_COLS:
        dupes = df.loc[:, df.columns == col]
        if dupes.shape[1] > 1:
            merged = dupes.bfill(axis=1).iloc[:, 0]
            df = df.loc[:, ~df.columns.duplicated()]
            df[col] = merged

    if "Latitude_decimal" not in df.columns or "Longitude_decimal" not in df.columns:
        return pd.DataFrame(), meta_folder, "Lat/Lon columns missing"

    df["Latitude_decimal"]  = pd.to_numeric(df["Latitude_decimal"],  errors="coerce")
    df["Longitude_decimal"] = pd.to_numeric(df["Longitude_decimal"], errors="coerce")
    df = df.dropna(subset=["Latitude_decimal", "Longitude_decimal"])
    df = df[
        df["Latitude_decimal"].between(-90, 90) &
        df["Longitude_decimal"].between(-180, 180)
    ]

    print(f"[{instrument_type}] Valid stations: {len(df)}")
    return df, meta_folder, None


def _df_to_stations(df: pd.DataFrame, instrument_type: str) -> list[dict]:
    stations = []
    for row in df.itertuples(index=False):
        try:
            raw_file  = getattr(row, "SourceFile", "")
            file_name = (
                str(raw_file)
                if pd.notna(raw_file) and str(raw_file).strip()
                else "N/A"
            )
            source_raw   = str(getattr(row, "SourceFolder", "N/A"))
            source_clean = source_raw.replace(".csv", "").replace("combined_metadata_", "")

            lat = float(row.Latitude_decimal)
            lon = float(row.Longitude_decimal)

            stations.append({
                "type":        instrument_type,
                "latitude":    lat,
                "longitude":   lon,
                "ship":        infer_ship_from_metadata(row._asdict()),
                "cruise":      str(getattr(row, "Cruise",          "N/A")),
                "station":     str(getattr(row, "Station Number",  "N/A")),
                "datetime":    str(getattr(row, "Datetime",        "N/A")),
                "depth":       str(getattr(row, "Station Depth",   "N/A")),
                "source":      source_clean,
                "file_name":   file_name,
                "folder_path": file_name,
                "in_eez":      is_in_eez(lat, lon),
            })
        except Exception as e:
            print(f"[{instrument_type}] Skipping row: {e}")
    return stations


# ==========================================
# STATION METADATA CACHE (SQLite-backed, mirrors profile cache pattern)
# ==========================================

def _meta_folder_mtime(folder: str) -> float:
    """Latest mtime across all metadata CSVs in folder."""
    if not os.path.isdir(folder):
        return 0.0
    files = [
        os.path.join(folder, f)
        for f in os.listdir(folder)
        if f.endswith(".csv")
    ]
    return max((os.path.getmtime(f) for f in files), default=0.0)


def _station_table(instrument_type: str) -> str:
    return f"stations_{instrument_type}"


STATION_COLUMNS = [
    "type", "latitude", "longitude", "ship", "cruise",
    "station", "datetime", "depth", "source", "file_name", "folder_path",
    "in_eez",
]


def _station_cache_is_current(instrument_type: str, meta_folder: str) -> bool:
    if not os.path.isfile(CACHE_DB):
        return False
    latest = _meta_folder_mtime(meta_folder)
    try:
        with sqlite3.connect(CACHE_DB) as conn:
            rows = dict(conn.execute("SELECT key, value FROM metadata").fetchall())
    except sqlite3.OperationalError:
        return False
    return (
        rows.get(_meta_key(instrument_type, "meta_source_path")) == meta_folder
        and float(rows.get(_meta_key(instrument_type, "meta_source_mtime"), -1)) == latest
    )


def _build_station_cache(instrument_type: str) -> list[dict]:
    """Parse metadata CSVs fresh, persist into SQLite, and return the station list."""
    df, meta_folder, error = _load_meta_df(instrument_type)
    if error:
        print(f"[{instrument_type}] Metadata load error: {error}")
        stations: list[dict] = []
    else:
        stations = _df_to_stations(df, instrument_type)

    table = _station_table(instrument_type)
    latest_mtime = _meta_folder_mtime(meta_folder)

    os.makedirs(CACHE_DIR, exist_ok=True)

    with sqlite3.connect(CACHE_DB) as conn:
        conn.execute("PRAGMA journal_mode=WAL")

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS metadata (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )
            """
        )

        conn.execute(f"DROP TABLE IF EXISTS {table}")
        conn.execute(
            f"""
            CREATE TABLE {table} (
                type TEXT, latitude REAL, longitude REAL,
                ship TEXT, cruise TEXT, station TEXT,
                datetime TEXT, depth TEXT, source TEXT,
                file_name TEXT, folder_path TEXT, in_eez INTEGER
            )
            """
        )

        if stations:
            conn.executemany(
                f"""
                INSERT INTO {table}
                ({", ".join(STATION_COLUMNS)})
                VALUES ({", ".join("?" for _ in STATION_COLUMNS)})
                """,
                [tuple(s.get(c) for c in STATION_COLUMNS) for s in stations],
            )

        conn.execute(f"CREATE INDEX IF NOT EXISTS idx_{table}_file ON {table}(file_name)")

        conn.execute(
            "INSERT OR REPLACE INTO metadata VALUES (?, ?)",
            (_meta_key(instrument_type, "meta_source_path"), meta_folder),
        )
        conn.execute(
            "INSERT OR REPLACE INTO metadata VALUES (?, ?)",
            (_meta_key(instrument_type, "meta_source_mtime"), str(latest_mtime)),
        )

    print(f"[{instrument_type}] Station cache built: {len(stations):,} rows.")
    return stations


def _load_station_cache_from_db(instrument_type: str) -> list[dict]:
    table = _station_table(instrument_type)
    with sqlite3.connect(CACHE_DB) as conn:
        conn.row_factory = sqlite3.Row
        try:
            rows = conn.execute(f"SELECT * FROM {table}").fetchall()
        except sqlite3.OperationalError:
            return []
    return [dict(row) for row in rows]


def ensure_station_cache(instrument_type: str, force: bool = False) -> list[dict]:
    """
    Returns the station list for instrument_type, using SQLite as the
    persistent cache. Only re-parses metadata CSVs when:
      - force=True (explicit reload requested), or
      - the metadata folder's contents have changed since last cache build.
    """
    meta_folder = _resolve_meta_folder(instrument_type)

    if not force and _station_cache_is_current(instrument_type, meta_folder):
        stations = _load_station_cache_from_db(instrument_type)
        print(
            f"[{instrument_type}] Station cache is current "
            f"({len(stations):,} rows) — skipped reparse."
        )
    else:
        stations = _build_station_cache(instrument_type)

    _station_cache[instrument_type] = stations
    return stations


def _station_is_eez_restricted(stem: str) -> bool:
    """Look up whether a given station stem falls inside the EEZ, using
    whichever instrument cache has it loaded."""
    for instrument_type in ["ctd", "xbt", "xctd"]:
        for s in _station_cache.get(instrument_type, []):
            fname = (s.get("file_name") or "").strip().rsplit(".", 1)[0].lower()
            if fname == stem:
                return bool(s.get("in_eez"))
    return False


# ==========================================
# UNIFIED "ALL STATIONS" CACHE — for initial map load
# ==========================================
# Combines ctd + xbt + xctd station metadata into a single table,
# pre-filtered to a fixed recent time window, so the map's first
# load is one fast query instead of three + client-side merging.

UNIFIED_TABLE = "stations_all"

INITIAL_LOAD_YEARS = int(os.environ.get("SEASNAP_INITIAL_LOAD_YEARS", 2))


def _initial_window_dates() -> tuple[date, date]:
    """Return (start_date, end_date) for the fixed initial-load window."""
    today = date.today()
    start = date(today.year - INITIAL_LOAD_YEARS, today.month, today.day)
    return start, today


def _unified_cache_key(kind: str) -> str:
    return f"all:{kind}"


def _unified_cache_is_current() -> bool:
    """
    Current if every instrument's underlying metadata folder is unchanged
    since the unified table was last built, AND the window hasn't changed.
    """
    if not os.path.isfile(CACHE_DB):
        return False

    try:
        with sqlite3.connect(CACHE_DB) as conn:
            rows = dict(conn.execute("SELECT key, value FROM metadata").fetchall())
    except sqlite3.OperationalError:
        return False

    stored_years = rows.get(_unified_cache_key("window_years"))
    if stored_years != str(INITIAL_LOAD_YEARS):
        return False

    for instrument_type in ["ctd", "xbt", "xctd"]:
        meta_folder = _resolve_meta_folder(instrument_type)
        latest = _meta_folder_mtime(meta_folder)
        stored_path  = rows.get(_meta_key(instrument_type, "meta_source_path"))
        stored_mtime = rows.get(_meta_key(instrument_type, "meta_source_mtime"))
        if stored_path != meta_folder or stored_mtime is None or float(stored_mtime) != latest:
            return False

    return True


def _build_unified_station_cache() -> list[dict]:
    """
    Ensures each instrument's per-type station cache is current, merges
    them, filters to the fixed initial-load window, and persists the
    result into a single stations_all table.
    """
    start_date, end_date = _initial_window_dates()

    all_stations: list[dict] = []
    for instrument_type in ["ctd", "xbt", "xctd"]:
        stations = ensure_station_cache(instrument_type)
        for s in stations:
            station_date = _parse_date(s.get("datetime", ""))
            if station_date is None or (start_date <= station_date <= end_date):
                all_stations.append(s)

    os.makedirs(CACHE_DIR, exist_ok=True)

    with sqlite3.connect(CACHE_DB) as conn:
        conn.execute("PRAGMA journal_mode=WAL")

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS metadata (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )
            """
        )

        conn.execute(f"DROP TABLE IF EXISTS {UNIFIED_TABLE}")
        conn.execute(
            f"""
            CREATE TABLE {UNIFIED_TABLE} (
                type TEXT, latitude REAL, longitude REAL,
                ship TEXT, cruise TEXT, station TEXT,
                datetime TEXT, depth TEXT, source TEXT,
                file_name TEXT, folder_path TEXT, in_eez INTEGER
            )
            """
        )

        if all_stations:
            conn.executemany(
                f"""
                INSERT INTO {UNIFIED_TABLE}
                ({", ".join(STATION_COLUMNS)})
                VALUES ({", ".join("?" for _ in STATION_COLUMNS)})
                """,
                [tuple(s.get(c) for c in STATION_COLUMNS) for s in all_stations],
            )

        conn.execute(f"CREATE INDEX IF NOT EXISTS idx_{UNIFIED_TABLE}_type ON {UNIFIED_TABLE}(type)")
        conn.execute(f"CREATE INDEX IF NOT EXISTS idx_{UNIFIED_TABLE}_latlon ON {UNIFIED_TABLE}(latitude, longitude)")

        conn.execute(
            "INSERT OR REPLACE INTO metadata VALUES (?, ?)",
            (_unified_cache_key("window_years"), str(INITIAL_LOAD_YEARS)),
        )

    print(
        f"[unified] Built stations_all: {len(all_stations):,} rows "
        f"(window: {start_date} to {end_date})"
    )
    return all_stations


def _load_unified_cache_from_db() -> list[dict]:
    with sqlite3.connect(CACHE_DB) as conn:
        conn.row_factory = sqlite3.Row
        try:
            rows = conn.execute(f"SELECT * FROM {UNIFIED_TABLE}").fetchall()
        except sqlite3.OperationalError:
            return []
    return [dict(row) for row in rows]


def ensure_unified_station_cache(force: bool = False) -> list[dict]:
    """
    Returns the merged ctd+xbt+xctd station list, filtered to the fixed
    initial-load window, using SQLite as the persistent cache.
    """
    if not force and _unified_cache_is_current():
        stations = _load_unified_cache_from_db()
        print(f"[unified] Cache is current ({len(stations):,} rows) — skipped rebuild.")
    else:
        stations = _build_unified_station_cache()

    return stations


class SpatialBox(BaseModel):
    latMin: float
    latMax: float
    lonMin: float
    lonMax: float
    dateFrom: Optional[str] = None   # "YYYY-MM-DD"
    dateTo:   Optional[str] = None   # "YYYY-MM-DD"


def _parse_date(raw: str | None):
    """Parse YYYY-MM-DD, DD-MM-YYYY, MM/DD/YYYY to a date object. Returns None on failure."""
    if not raw or raw == "N/A":
        return None
    m = re.match(r'^(\d{4})-(\d{2})-(\d{2})', raw)
    if m:
        try: return date(int(m[1]), int(m[2]), int(m[3]))
        except: return None
    m = re.match(r'^(\d{2})-(\d{2})-(\d{4})', raw)
    if m:
        try: return date(int(m[3]), int(m[2]), int(m[1]))
        except: return None
    m = re.match(r'^(\d{2})/(\d{2})/(\d{4})', raw)
    if m:
        try: return date(int(m[3]), int(m[1]), int(m[2]))
        except: return None
    return None


def _stations_in_box(box: SpatialBox) -> list[dict]:
    date_from = _parse_date(box.dateFrom)
    date_to   = _parse_date(box.dateTo)

    stations = []
    for instrument_type in ["ctd", "xbt", "xctd"]:
        for station in _station_cache.get(instrument_type, []):
            lat = float(station["latitude"])
            lon = float(station["longitude"])

            if not (box.latMin <= lat <= box.latMax and
                    box.lonMin <= lon <= box.lonMax):
                continue

            if date_from or date_to:
                station_date = _parse_date(station.get("datetime", ""))
                if station_date:
                    if date_from and station_date < date_from:
                        continue
                    if date_to and station_date > date_to:
                        continue

            stations.append(station)

    return stations


# ==========================================
# DATA REQUISITION WORKFLOW
# ==========================================

SMTP_HOST = os.environ.get("SEASNAP_SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.environ.get("SEASNAP_SMTP_PORT", 587))
SMTP_USER = os.environ.get("SEASNAP_SMTP_USER")
SMTP_PASS = os.environ.get("SEASNAP_SMTP_PASS")
TEAM_EMAIL = os.environ.get("SEASNAP_TEAM_EMAIL", "data-requests@incois.gov.in")

REQUISITION_DIR = os.path.join(CACHE_DIR, "requisitions")

TEMPLATE_PATH = os.environ.get(
    "SEASNAP_REQUISITION_TEMPLATE",
    os.path.join(BASE_DIR, "templates", "Data_Requisition_Template.docx"),
)


def _replace_placeholder(paragraph, token: str, value: str) -> None:
    """Replace a {{TOKEN}} placeholder that lives in a paragraph, even if
    Word split it across multiple runs. Our template places each token in
    its own isolated run, so this is safe and doesn't disturb other text
    in the same paragraph."""
    if token not in paragraph.text:
        return
    full_text = paragraph.text.replace(token, value)
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)


def _replace_all(doc, replacements: dict[str, str]) -> None:
    def _paragraphs_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    yield from _paragraphs_in_table(nested)

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        all_paragraphs.extend(_paragraphs_in_table(table))

    for paragraph in all_paragraphs:
        for token, value in replacements.items():
            _replace_placeholder(paragraph, token, value)


def _convert_docx_to_pdf(docx_path: str, out_dir: str) -> str:
    """Requires LibreOffice (`soffice`) installed on the server."""
    import subprocess
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"soffice conversion failed: {result.stderr.strip()}")
    pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.isfile(pdf_path):
        raise RuntimeError("soffice did not produce the expected PDF output")
    return pdf_path


def fill_requisition_pdf(form: "RequisitionForm", req_id: str) -> str:
    """Autofills the official Data Requisition Form template with this
    request's details and returns the path to the generated PDF."""
    from docx import Document as DocxDocument

    if not os.path.isfile(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Requisition template not found at {TEMPLATE_PATH}. "
            f"Set SEASNAP_REQUISITION_TEMPLATE."
        )

    doc = DocxDocument(TEMPLATE_PATH)

    is_consultancy = form.request_type == "consultancy"
    project_and_cost = form.purpose
    if form.project_cost:
        project_and_cost += f" (Estimated cost: {form.project_cost})"

    replacements = {
        "{{INSTITUTION_ADDRESS}}":        form.institution_address or form.organization,
        "{{OFFICER_NAME_DESIGNATION}}":   f"{form.name}, {form.officer_designation}".rstrip(", "),
        "{{PARAMETERS}}":                 form.parameters or "N/A",
        "{{PLATFORM_INSTRUMENT}}":        f"{form.instrument_type.upper()} — {form.station_file}",
        "{{PERIOD}}":                     form.period or "N/A",
        "{{PROJECT_AND_COST}}":           project_and_cost,
        "{{CHECKBOX_OWN_RESEARCH}}":      "☑" if not is_consultancy else "☐",
        "{{CHECKBOX_CONSULTANCY}}":       "☑" if is_consultancy else "☐",
        "{{GOVT_APPROVAL_DETAILS}}":      (form.govt_approval_details or "N/A") if is_consultancy else "N/A",
        "{{OFFICER_SIGNATURE}}":          form.name,  # typed name; physical/formal sign-off happens at approval
        "{{STATION_PLACE}}":              form.organization,
        "{{DATE}}":                       str(date.today()),
    }

    _replace_all(doc, replacements)

    os.makedirs(REQUISITION_DIR, exist_ok=True)
    filled_docx_path = os.path.join(REQUISITION_DIR, f"{req_id}_filled.docx")
    doc.save(filled_docx_path)

    pdf_path = _convert_docx_to_pdf(filled_docx_path, REQUISITION_DIR)
    return pdf_path




def _init_requisition_table():
    with sqlite3.connect(CACHE_DB) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS requisitions (
                id TEXT PRIMARY KEY,
                name TEXT, email TEXT, organization TEXT,
                purpose TEXT, station_file TEXT, instrument_type TEXT,
                status TEXT DEFAULT 'pending',
                created_at TEXT
            )
        """)


def _send_email(to_addr: str, subject: str, body: str, attachment_path: str = None, attachment_name: str = None):
    msg = EmailMessage()
    msg["From"] = SMTP_USER
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.set_content(body)

    if attachment_path:
        with open(attachment_path, "rb") as f:
            msg.add_attachment(
                f.read(), maintype="application", subtype="pdf",
                filename=attachment_name or os.path.basename(attachment_path),
            )

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
        server.starttls()
        server.login(SMTP_USER, SMTP_PASS)
        server.send_message(msg)


class RequisitionForm(BaseModel):
    name: str
    email: str
    organization: str
    purpose: str
    station_file: str
    instrument_type: str

    # Fields matching the official Data Requisition Form. Optional with
    # sensible defaults so older frontend builds don't break, but the
    # updated RequisitionForm.jsx should collect all of these.
    institution_address: str = ""
    officer_designation: str = ""
    parameters: str = ""
    period: str = ""
    project_cost: Optional[str] = None
    request_type: str = "own_research"   # "own_research" | "consultancy"
    govt_approval_details: Optional[str] = None


# ==========================================
# API ENDPOINTS
# ==========================================

def _validate_type(instrument_type: str) -> None:
    if instrument_type not in INSTRUMENT_CONFIG:
        raise HTTPException(status_code=400, detail=f"Unknown instrument type: '{instrument_type}'")


@app.post("/load-meta", dependencies=[Depends(require_api_key)])
def load_meta(type: str = Query("ctd"), force: bool = Query(False)):
    """
    Loads station metadata for the given instrument type.

    By default this is a no-op re-parse: if the metadata folder hasn't
    changed since the last build, stations are served from the SQLite
    cache directly. Pass force=true to always re-parse from CSV.
    """
    instrument_type = type.lower()
    _validate_type(instrument_type)

    stations = ensure_station_cache(instrument_type, force=force)

    if not stations:
        raise HTTPException(status_code=500, detail="No valid stations loaded")

    return {
        "type":    instrument_type,
        "count":   len(stations),
        "message": "Metadata loaded successfully",
    }


@app.get("/stations", dependencies=[Depends(require_api_key)])
def get_stations(type: str = Query("ctd")):
    instrument_type = type.lower()
    _validate_type(instrument_type)

    if instrument_type not in _station_cache:
        ensure_station_cache(instrument_type)

    return {"stations": _station_cache.get(instrument_type, [])}


@app.get("/stations/initial", dependencies=[Depends(require_api_key)])
def get_initial_stations(force: bool = Query(False)):
    """
    Returns merged station metadata (CTD + XBT + XCTD) for the fixed
    lookback window, from a single unified table. Intended for the
    initial map load — one fast query instead of three separate calls.
    """
    stations = ensure_unified_station_cache(force=force)
    return {
        "stations": stations,
        "count": len(stations),
        "window_years": INITIAL_LOAD_YEARS,
    }


@app.get("/profile/{station_file:path}", dependencies=[Depends(require_api_key)])
def get_profile(station_file: str, type: str = Query(None)):
    station_file = station_file.strip()
    stem = re.sub(r"_metadata\.csv$", "", station_file, flags=re.IGNORECASE)
    stem = stem.rsplit(".", 1)[0].strip().lower()

    # Server-side EEZ enforcement — this is the actual security boundary.
    # No amount of frontend hiding matters if this check isn't here: the
    # raw JSON never leaves the server for a restricted station at all.
    if _station_is_eez_restricted(stem):
        raise HTTPException(
            status_code=403,
            detail="This station is inside a restricted EEZ. Submit a data requisition to request access.",
        )

    if type and type.lower() in INSTRUMENT_CONFIG:
        search_types = [type.lower()]
    else:
        search_types = ["ctd", "xbt", "xctd"]

    for instrument_type in search_types:
        cfg      = INSTRUMENT_CONFIG[instrument_type]
        table    = cfg["table"]
        out_cols = cfg["output_columns"]

        folder = _resolve_folder(instrument_type)

        if not _cache_is_current(instrument_type, folder):
            _build_cache(instrument_type, folder)

        col_list = ", ".join(f'"{c}"' for c in out_cols)
        query    = f'SELECT {col_list} FROM {table} WHERE stem = ? ORDER BY depSM'

        try:
            with sqlite3.connect(CACHE_DB) as conn:
                conn.execute("PRAGMA cache_size=-64000")
                conn.execute("PRAGMA temp_store=MEMORY")
                conn.row_factory = sqlite3.Row
                rows = conn.execute(query, (stem,)).fetchall()
        except sqlite3.OperationalError:
            rows = []

        if rows:
            print(f"FOUND PROFILE: stem={stem} table={table} type={instrument_type}")
            out_col_set = set(out_cols)
            return [
                {
                    **{col: (row[col] if col in out_col_set else None) for col in ALL_OUTPUT_COLUMNS},
                    "instrument_type": instrument_type,
                }
                for row in rows
            ]

    raise HTTPException(
        status_code=404,
        detail=f"No profile found for '{stem}' in instrument type '{type or 'any'}'"
    )


@app.post("/spatial-profile", dependencies=[Depends(require_api_key)])
def get_spatial_profile(box: SpatialBox):

    selected_stations = _stations_in_box(box)

    # Drop EEZ-restricted stations before any data is queried — restricted
    # rows should never be pulled into merged_rows in the first place.
    restricted_count = sum(1 for s in selected_stations if s.get("in_eez"))
    selected_stations = [s for s in selected_stations if not s.get("in_eez")]

    if not selected_stations:
        return {
            "mode": "spatial",
            "station_count": 0,
            "row_count": 0,
            "restricted_count": restricted_count,
            "data": [],
        }

    merged_rows = []
    counts = {"ctd": 0, "xbt": 0, "xctd": 0}

    with sqlite3.connect(CACHE_DB) as conn:
        conn.execute("PRAGMA cache_size=-64000")
        conn.execute("PRAGMA temp_store=MEMORY")
        conn.row_factory = sqlite3.Row

        for instrument_type in ["ctd", "xbt", "xctd"]:
            type_stations = [s for s in selected_stations if s["type"] == instrument_type]
            if not type_stations:
                continue

            cfg      = INSTRUMENT_CONFIG[instrument_type]
            table    = cfg["table"]
            out_cols = cfg["output_columns"]
            out_col_set = set(out_cols)

            stems = [
                s["file_name"].strip().rsplit(".", 1)[0].lower()
                for s in type_stations
            ]

            placeholders = ",".join("?" * len(stems))
            col_list     = ", ".join(f'"{c}"' for c in out_cols)
            query = (
                f'SELECT stem, {col_list} FROM {table} '
                f'WHERE stem IN ({placeholders}) '
                f'ORDER BY stem, depSM'
            )

            rows = conn.execute(query, stems).fetchall()
            counts[instrument_type] = len(set(r["stem"] for r in rows))

            for row in rows:
                merged_rows.append({
                    **{col: (row[col] if col in out_col_set else None)
                       for col in ALL_OUTPUT_COLUMNS},
                    "instrument_type": instrument_type,
                    "station_file":    row["stem"],
                })

    return {
        "mode": "spatial",
        "station_count": len(selected_stations),
        "ctd_count": counts["ctd"],
        "xbt_count": counts["xbt"],
        "xctd_count": counts["xctd"],
        "row_count": len(merged_rows),
        "restricted_count": restricted_count,
        "data": merged_rows,
    }


@app.post("/requisition/submit", dependencies=[Depends(require_api_key)])
def submit_requisition(form: RequisitionForm):
    """
    Records the request, autofills the official Data Requisition Form
    with this request's details, converts it to PDF, and emails it to
    the internal team for review. The team signs off outside this
    system, then uses /requisition/{id}/approve to send the signed
    document back to the requester. The filled PDF is also made
    available to the requester immediately via request_pdf_url.
    """
    _init_requisition_table()
    req_id = str(uuid.uuid4())

    # Autofill the actual INCOIS Data Requisition Form template and
    # convert to PDF. Wrapped so a template/LibreOffice failure doesn't
    # crash the whole request with a 500 — which the browser then
    # misreports as a CORS error, since FastAPI's CORSMiddleware can't
    # attach its headers to an unhandled-exception response.
    pdf_path_used = None
    try:
        pdf_path_used = fill_requisition_pdf(form, req_id)
    except Exception as e:
        print(f"[requisition] Form autofill/PDF generation failed: {e}")

    with sqlite3.connect(CACHE_DB) as conn:
        conn.execute(
            "INSERT INTO requisitions VALUES (?, ?, ?, ?, ?, ?, ?, 'pending', ?)",
            (req_id, form.name, form.email, form.organization, form.purpose,
             form.station_file, form.instrument_type, str(date.today())),
        )

    try:
        _send_email(
            TEAM_EMAIL,
            subject=f"New data requisition — {form.station_file}",
            body=(
                f"New EEZ data request from {form.name} ({form.email}, {form.organization}).\n\n"
                f"Purpose: {form.purpose}\n"
                f"Station: {form.station_file} ({form.instrument_type})\n\n"
                f"To approve and send the signed document back to the requester, "
                f"use request ID: {req_id}"
            ),
            attachment_path=pdf_path_used,
            attachment_name="requisition_request.pdf",
        )
    except Exception as e:
        print(f"[requisition] Email to team failed: {e}")

    # Also send the requester their own copy of the filled (unsigned) form
    # right away, so they have a record of what was submitted — separate
    # from the signed copy they'll receive later via /approve.
    try:
        _send_email(
            form.email,
            subject=f"Your SeaSnap data requisition — {form.station_file}",
            body=(
                f"Hi {form.name},\n\n"
                f"Thanks for submitting a data requisition for station "
                f"{form.station_file} ({form.instrument_type}). We've attached "
                f"a copy of your filled request form.\n\n"
                f"Our team will review it and, once approved and signed, send "
                f"the authorized data directly to this email address.\n\n"
                f"Request ID: {req_id}"
            ),
            attachment_path=pdf_path_used,
            attachment_name="requisition_request.pdf",
        )
    except Exception as e:
        print(f"[requisition] Email to requester failed: {e}")

    return {
        "request_id": req_id,
        "status": "pending",
        "message": "Request submitted. The team has been notified.",
        "request_pdf_url": f"/requisition/{req_id}/request-pdf" if pdf_path_used else None,
    }


@app.get("/requisition/{request_id}/request-pdf", dependencies=[Depends(require_api_key)])
def get_requisition_pdf(request_id: str):
    """Serves the autofilled request PDF back to the requester (or anyone
    holding the request_id + API key) right after submission."""
    from fastapi.responses import FileResponse
    pdf_path = os.path.join(REQUISITION_DIR, f"{request_id}_filled.pdf")
    if not os.path.isfile(pdf_path):
        raise HTTPException(status_code=404, detail="Filled PDF not found for this request.")
    return FileResponse(pdf_path, media_type="application/pdf", filename=f"{request_id}_requisition.pdf")


@app.post("/requisition/{request_id}/approve")
def approve_requisition(request_id: str, signed_doc: UploadFile = File(...)):
    """
    Internal-team-only endpoint: upload the signed PDF once review is
    complete. Emails it directly to the original requester.

    NOTE: intentionally NOT gated by the same public API key — this is
    meant to be used by internal staff only. Give this its own separate
    credential before deploying, rather than reusing the frontend's key.
    """
    with sqlite3.connect(CACHE_DB) as conn:
        conn.row_factory = sqlite3.Row
        row = conn.execute("SELECT * FROM requisitions WHERE id = ?", (request_id,)).fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Request not found")

    signed_path = os.path.join(REQUISITION_DIR, f"{request_id}_signed.pdf")
    with open(signed_path, "wb") as f:
        f.write(signed_doc.file.read())

    _send_email(
        row["email"],
        subject="Your SeaSnap data requisition has been approved",
        body=(
            f"Hi {row['name']},\n\n"
            f"Your request for station {row['station_file']} has been reviewed and "
            f"approved. Please find the signed authorization attached."
        ),
        attachment_path=signed_path,
        attachment_name="signed_authorization.pdf",
    )

    with sqlite3.connect(CACHE_DB) as conn:
        conn.execute("UPDATE requisitions SET status = 'approved' WHERE id = ?", (request_id,))

    return {"status": "approved", "message": "Signed document sent to requester."}


@app.on_event("startup")
async def startup():
    import threading

    def _warm():
        if os.path.isfile(CACHE_DB):
            with sqlite3.connect(CACHE_DB) as conn:
                conn.execute("PRAGMA wal_checkpoint(PASSIVE)")

        for t in ["ctd", "xbt", "xctd"]:
            try:
                ensure_cache(t)
                ensure_station_cache(t)
                print(f"[startup] {t} warmed up")
            except Exception as e:
                print(f"[startup] {t} failed: {e}")

        try:
            ensure_unified_station_cache()
            print("[startup] unified stations_all warmed up")
        except Exception as e:
            print(f"[startup] unified stations_all failed: {e}")

    threading.Thread(target=_warm, daemon=True).start()